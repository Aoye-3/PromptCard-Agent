from docx import Document
import os
import re


def table_to_markdown(table_data):
    if not table_data:
        return ''
    
    md_lines = []
    md_lines.append('| ' + ' | '.join(table_data[0]) + ' |')
    md_lines.append('| ' + ' | '.join(['---'] * len(table_data[0])) + ' |')
    for row in table_data[1:]:
        md_lines.append('| ' + ' | '.join(row) + ' |')
    
    return '\n'.join(md_lines)


def main():
    docx_path = r'f:\.workSpace\IICL-CardInterface\promptcard-v4\104种教学版+提示词.docx'
    output_dir = r'f:\.workSpace\IICL-CardInterface\promptcard-v4\104种教学提示词'
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    doc = Document(docx_path)
    
    md_lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        level = 0
        if para.style.name.startswith('Heading'):
            try:
                level = int(para.style.name.replace('Heading ', ''))
            except:
                pass
        
        if level == 1:
            md_lines.append(f'# {text}')
        elif level == 2:
            md_lines.append(f'## {text}')
        elif level == 3:
            md_lines.append(f'### {text}')
        elif level == 4:
            md_lines.append(f'#### {text}')
        else:
            md_lines.append(text)
    
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        
        if table_data:
            md_lines.append('\n' + table_to_markdown(table_data))
    
    full_md = '\n'.join(md_lines)
    
    with open(os.path.join(output_dir, '完整文档.md'), 'w', encoding='utf-8') as f:
        f.write(full_md)
    
    print(f'已保存完整文档到 {output_dir}')
    print('现在开始按分类拆分...')
    
    lines = full_md.split('\n')
    categories = {}
    current_title = None
    current_content = []
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('#'):
            if current_title and current_content:
                categories[current_title] = current_content
            
            current_title = line.lstrip('#').strip()
            current_content = [line]
        else:
            if current_title:
                current_content.append(line)
    
    if current_title and current_content:
        categories[current_title] = current_content
    
    for idx, (title, content) in enumerate(categories.items()):
        safe_name = re.sub(r'[\\/:*?"<>|]', '_', title)
        safe_name = safe_name[:50]
        filename = f'{idx+1:02d}_{safe_name}.md'
        
        with open(os.path.join(output_dir, filename), 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
    
    print(f'已按分类拆分为 {len(categories)} 个文件')


if __name__ == '__main__':
    main()
